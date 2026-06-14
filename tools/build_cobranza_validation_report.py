from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "informe-validacion-modulo-cobranza.docx"
PDF_PATH = OUT / "informe-validacion-modulo-cobranza.pdf"
FLOW_IMG = OUT / "flujo-modulo-cobranza-validacion.png"


def get_font(size=28, bold=False):
    try:
        return ImageFont.truetype("arialbd.ttf" if bold else "arial.ttf", size)
    except Exception:
        return ImageFont.load_default()


def draw_box(draw, xy, text, fill="#F4F6F9", outline="#2E74B5"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    words = text.split()
    lines, line = [], ""
    for word in words:
        if len((line + " " + word).strip()) > 19:
            lines.append(line)
            line = word
        else:
            line = (line + " " + word).strip()
    if line:
        lines.append(line)
    font = get_font(22, True)
    y = y1 + ((y2 - y1) - len(lines) * 28) / 2
    for ln in lines:
        bbox = draw.textbbox((0, 0), ln, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), ln, font=font, fill="#142030")
        y += 28


def make_flow():
    img = Image.new("RGB", (1400, 620), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 45), "Flujo mensual del modulo de cobranza", font=get_font(42, True), fill="#142030")
    d.text((60, 100), "Importacion, asignacion, cobro, control y rendicion", font=get_font(25), fill="#4B5870")
    steps = [
        "Importar listado",
        "Detectar dependencia",
        "Asignar cobrador",
        "Cargar cobros",
        "Controlar alertas",
        "Rendir cobranza",
    ]
    x = 55
    for i, step in enumerate(steps):
        fill = "#E8EEF5" if i % 2 == 0 else "#F4F6F9"
        draw_box(d, (x, 230, x + 195, 360), step, fill=fill)
        if i < len(steps) - 1:
            d.line((x + 195, 295, x + 232, 295), fill="#2E74B5", width=5)
            d.polygon([(x + 232, 295), (x + 218, 285), (x + 218, 305)], fill="#2E74B5")
        x += 220
    d.text((60, 450), "La app conserva asignaciones, observaciones, deuda por tickets y control de talonarios para los meses siguientes.", font=get_font(25), fill="#4B5870")
    img.save(FLOW_IMG)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def doc_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell_text(table.rows[0].cells[i], header, True)
        shade(table.rows[0].cells[i], "F4F6F9")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def build_docx():
    doc = Document()
    configure_doc(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Informe para validacion - Modulo de Cobranza")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(20, 32, 48)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema de gestion para cartera de afiliados, tickets, recibos y rendicion")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(31, 77, 120)
    doc.add_picture(str(FLOW_IMG), width=Inches(6.5))

    doc.add_heading("1. Objetivo del modulo", level=1)
    doc.add_paragraph(
        "El modulo de cobranza tiene como objetivo administrar la carga mensual de tickets, organizar la cartera por "
        "dependencia y cobrador, registrar cobros por tickets y recibos, controlar observaciones pendientes, detectar "
        "polizas con atraso y cerrar la rendicion con mayor seguridad administrativa."
    )

    doc.add_heading("2. Importacion mensual y dependencias", level=1)
    bullets(doc, [
        "El listado mensual se recibira en Excel con los afiliados y tickets de cobranza.",
        "La app debera detectar automaticamente la dependencia desde el codigo del plan.",
        "Ejemplo: A CA238 se interpreta como plan A y dependencia 238.",
        "Ejemplo: G CG238 se interpreta como plan G y dependencia 238.",
        "La regla debe ser flexible porque existen mas dependencias y otros tipos de planes.",
    ])
    doc_table(doc, ["Dato original", "Interpretacion esperada"], [
        ["A CA238", "Plan A, dependencia 238"],
        ["G CG238", "Plan G, dependencia 238"],
        ["A CA269", "Plan A, dependencia 269"],
        ["G CG269", "Plan G, dependencia 269"],
        ["Otros planes", "Detectar segun regla disponible o marcar para revision"],
    ], widths=[2.0, 4.5])

    doc.add_heading("3. Asignacion de cobradores", level=1)
    bullets(doc, [
        "Primero los tickets se separan por dependencia.",
        "Luego, dentro de cada dependencia, pueden existir dos o mas cobradores.",
        "Cada dependencia tendra un cobrador por defecto llamado OFICINA.",
        "Las polizas nuevas se asignan automaticamente a OFICINA.",
        "Si administracion cambia el cobrador de una poliza, ese cambio queda guardado para el mes siguiente.",
        "Si no se cambia, la poliza sigue asignada a OFICINA.",
    ])
    doc_table(doc, ["Caso", "Accion automatica"], [
        ["Poliza nueva", "Se carga en su dependencia y queda asignada al cobrador OFICINA."],
        ["Poliza existente", "Conserva el cobrador que tenia guardado del mes anterior."],
        ["Cambio manual de cobrador", "Queda registrado como asignacion para futuras importaciones."],
        ["Cambio de dependencia", "Debe marcarse para revision administrativa."],
    ], widths=[2.0, 4.5])

    doc.add_heading("4. Observaciones obligatorias del ticket", level=1)
    doc.add_paragraph(
        "Al cargar o revisar tickets, administracion podra agregar observaciones cuando falten datos o sea necesario "
        "confirmar informacion del afiliado. Si el ticket tiene una observacion pendiente, el cobrador debera resolverla "
        "antes de registrar el cobro."
    )
    bullets(doc, [
        "Ejemplos: falta telefono, falta domicilio, confirmar titular, domicilio incompleto, dato dudoso.",
        "Si el cobrador pudo cobrar, se entiende que tuvo comunicacion con el asegurado.",
        "Por eso, la app no debe permitir registrar el ticket como cobrado sin completar la informacion o justificar por que no pudo hacerlo.",
        "Si el ticket no se cobra, el cobrador puede cargar una novedad libre sin obligacion de completar el dato.",
    ])

    doc.add_heading("5. Control de 4 tickets sin cobrar", level=1)
    bullets(doc, [
        "En general, el sistema emite tickets hasta cuatro meses de atraso.",
        "Cuando una poliza acumula 4 tickets sin cobrar, al mes siguiente puede dejar de salir en el listado.",
        "La poliza no necesariamente esta caida; puede cobrarse con recibo manual.",
        "La app debera guardar las polizas con 4 tickets pendientes.",
        "Al importar el siguiente mes, si una de esas polizas ya no aparece, se genera una alerta.",
    ])
    doc_table(doc, ["Alerta", "Accion sugerida"], [
        ["Poliza con 4 tickets pendientes", "Marcar como riesgo de no emision."],
        ["No aparece en el nuevo listado", "Alertar: cobrar con recibo manual."],
        ["Regularizacion", "Registrar recibo manual, novedad o estado de seguimiento."],
    ], widths=[2.4, 4.1])

    doc.add_heading("6. Talonarios de recibos", level=1)
    doc.add_paragraph(
        "Para mejorar el control de recibos manuales, los talonarios deberan precargarse con numeracion desde y hasta. "
        "Cuando se entreguen a un cobrador, se asignaran a su cobranza."
    )
    bullets(doc, [
        "Precargar talonario, rango desde/hasta y cantidad de recibos.",
        "Estado inicial: disponible.",
        "Al entregar, asignar a cobrador, dependencia y mes de cobranza.",
        "Validar que los recibos cargados pertenezcan al talonario asignado.",
        "Controlar duplicados, saltos de numeracion, anulados y recibos faltantes.",
        "Al rendir, revisar usados, pendientes, anulados y monto total.",
    ])
    doc_table(doc, ["Dato del talonario", "Uso"], [
        ["Numero de talonario", "Identifica el lote de recibos."],
        ["Desde / hasta", "Define el rango valido de numeracion."],
        ["Cobrador asignado", "Permite controlar quien tiene el talonario."],
        ["Dependencia y mes", "Relaciona el talonario con una cobranza concreta."],
        ["Estado", "Disponible, asignado, en uso, rendido, anulado o extraviado."],
    ], widths=[2.0, 4.5])

    doc.add_heading("7. Reportes y controles esperados", level=1)
    bullets(doc, [
        "Totales por dependencia.",
        "Totales por cobrador.",
        "Tickets recibidos, cobrados y pendientes.",
        "Polizas nuevas asignadas a OFICINA.",
        "Tickets con observaciones pendientes.",
        "Polizas con 4 tickets sin cobrar o sin emision en el mes siguiente.",
        "Recibos usados, pendientes, anulados y faltantes por talonario.",
        "Rendicion por cobrador y por dependencia.",
    ])

    doc.add_heading("8. Cierre para validacion", level=1)
    doc.add_paragraph(
        "Este informe resume las reglas funcionales propuestas para el modulo de cobranza. La validacion deberia confirmar "
        "si la deteccion de dependencias, la asignacion automatica a OFICINA, las observaciones obligatorias, el control "
        "de atraso y la gestion de talonarios coinciden con la forma real de trabajo."
    )

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="TitleX", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=23, leading=28,
        textColor=colors.HexColor("#142030"), alignment=TA_CENTER, spaceAfter=10
    ))
    styles.add(ParagraphStyle(
        name="SubtitleX", parent=styles["BodyText"], fontName="Helvetica", fontSize=12, leading=16,
        textColor=colors.HexColor("#1F4D78"), alignment=TA_CENTER, spaceAfter=14
    ))
    styles.add(ParagraphStyle(
        name="H1X", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=18,
        textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=7
    ))
    styles.add(ParagraphStyle(
        name="BodyX", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13,
        textColor=colors.HexColor("#202A3A"), spaceAfter=6
    ))
    styles.add(ParagraphStyle(
        name="BulletX", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12.5,
        leftIndent=14, firstLineIndent=-8, spaceAfter=4, textColor=colors.HexColor("#202A3A")
    ))
    styles.add(ParagraphStyle(
        name="SmallX", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.3, leading=10.5,
        textColor=colors.HexColor("#4B5870")
    ))
    return styles


def para(text, style):
    return Paragraph(text, style)


def pdf_bullets(story, styles, items):
    for item in items:
        story.append(para(f"• {item}", styles["BulletX"]))


def pdf_table(data, widths, styles):
    table = Table([[para(str(c), styles["SmallX"]) for c in row] for row in data],
                  colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, rightMargin=0.65 * inch, leftMargin=0.65 * inch,
                            topMargin=0.55 * inch, bottomMargin=0.55 * inch)
    story = [
        para("Informe para validacion - Modulo de Cobranza", styles["TitleX"]),
        para("Sistema de gestion para cartera de afiliados, tickets, recibos y rendicion", styles["SubtitleX"]),
        PdfImage(str(FLOW_IMG), width=6.6 * inch, height=2.92 * inch),
        Spacer(1, 0.12 * inch),
        para("1. Objetivo del modulo", styles["H1X"]),
        para("El modulo de cobranza tiene como objetivo administrar la carga mensual de tickets, organizar la cartera por dependencia y cobrador, registrar cobros por tickets y recibos, controlar observaciones pendientes, detectar polizas con atraso y cerrar la rendicion con mayor seguridad administrativa.", styles["BodyX"]),
        para("2. Importacion mensual y dependencias", styles["H1X"]),
    ]
    pdf_bullets(story, styles, [
        "El listado mensual se recibira en Excel con afiliados y tickets de cobranza.",
        "La app detectara automaticamente la dependencia desde el codigo del plan.",
        "Ejemplo: A CA238 significa plan A y dependencia 238.",
        "Ejemplo: G CG238 significa plan G y dependencia 238.",
        "La regla debe ser flexible para mas dependencias y otros planes.",
    ])
    story.append(pdf_table([
        ["Dato original", "Interpretacion esperada"],
        ["A CA238", "Plan A, dependencia 238"],
        ["G CG238", "Plan G, dependencia 238"],
        ["A CA269", "Plan A, dependencia 269"],
        ["G CG269", "Plan G, dependencia 269"],
        ["Otros planes", "Detectar segun regla disponible o marcar para revision"],
    ], [2.0, 4.45], styles))
    story.append(PageBreak())
    story.append(para("3. Asignacion de cobradores", styles["H1X"]))
    pdf_bullets(story, styles, [
        "Primero los tickets se separan por dependencia.",
        "Dentro de cada dependencia pueden existir dos o mas cobradores.",
        "Cada dependencia tendra un cobrador por defecto llamado OFICINA.",
        "Las polizas nuevas se asignan automaticamente a OFICINA.",
        "Si administracion cambia el cobrador, el cambio queda guardado para el mes siguiente.",
    ])
    story.append(pdf_table([
        ["Caso", "Accion automatica"],
        ["Poliza nueva", "Se carga en su dependencia y queda asignada al cobrador OFICINA."],
        ["Poliza existente", "Conserva el cobrador que tenia guardado del mes anterior."],
        ["Cambio manual", "Queda registrado como asignacion para futuras importaciones."],
        ["Cambio de dependencia", "Debe marcarse para revision administrativa."],
    ], [2.0, 4.45], styles))
    story.append(para("4. Observaciones obligatorias del ticket", styles["H1X"]))
    pdf_bullets(story, styles, [
        "Administracion podra agregar observaciones cuando falten datos o sea necesario confirmar informacion.",
        "Si el ticket tiene observacion pendiente, el cobrador debera resolverla antes de registrar el cobro.",
        "Si pudo cobrar, se entiende que tuvo comunicacion con el asegurado.",
        "La app no debe permitir registrar como cobrado sin completar el dato o justificar por que no pudo hacerlo.",
        "Si el ticket no se cobra, puede cargar novedad libre sin obligacion de completar el dato.",
    ])
    story.append(para("5. Control de 4 tickets sin cobrar", styles["H1X"]))
    pdf_bullets(story, styles, [
        "El sistema emite tickets hasta cuatro meses de atraso.",
        "Si una poliza acumula 4 tickets sin cobrar, al mes siguiente puede dejar de salir en el listado.",
        "La poliza no necesariamente esta caida; puede cobrarse con recibo manual.",
        "La app guardara esas polizas y alertara si no aparecen en el siguiente mes.",
    ])
    story.append(pdf_table([
        ["Alerta", "Accion sugerida"],
        ["Poliza con 4 tickets pendientes", "Marcar como riesgo de no emision."],
        ["No aparece en nuevo listado", "Alertar: cobrar con recibo manual."],
        ["Regularizacion", "Registrar recibo manual, novedad o seguimiento."],
    ], [2.4, 4.05], styles))
    story.append(PageBreak())
    story.append(para("6. Talonarios de recibos", styles["H1X"]))
    pdf_bullets(story, styles, [
        "Precargar talonarios con numeracion desde/hasta.",
        "Estado inicial: disponible.",
        "Al entregar, asignar a cobrador, dependencia y mes de cobranza.",
        "Validar que los recibos cargados pertenezcan al talonario asignado.",
        "Controlar duplicados, saltos, anulados, faltantes y pendientes.",
        "Al rendir, revisar usados, pendientes, anulados y monto total.",
    ])
    story.append(pdf_table([
        ["Dato del talonario", "Uso"],
        ["Numero de talonario", "Identifica el lote de recibos."],
        ["Desde / hasta", "Define el rango valido de numeracion."],
        ["Cobrador asignado", "Permite controlar quien tiene el talonario."],
        ["Dependencia y mes", "Relaciona el talonario con una cobranza concreta."],
        ["Estado", "Disponible, asignado, en uso, rendido, anulado o extraviado."],
    ], [2.0, 4.45], styles))
    story.append(para("7. Reportes y controles esperados", styles["H1X"]))
    pdf_bullets(story, styles, [
        "Totales por dependencia y por cobrador.",
        "Tickets recibidos, cobrados y pendientes.",
        "Polizas nuevas asignadas a OFICINA.",
        "Tickets con observaciones pendientes.",
        "Polizas con 4 tickets sin cobrar o sin emision en el mes siguiente.",
        "Recibos usados, pendientes, anulados y faltantes por talonario.",
        "Rendicion por cobrador y por dependencia.",
    ])
    story.append(para("8. Cierre para validacion", styles["H1X"]))
    story.append(para("Este informe resume las reglas funcionales propuestas para el modulo de cobranza. La validacion deberia confirmar si la deteccion de dependencias, la asignacion automatica a OFICINA, las observaciones obligatorias, el control de atraso y la gestion de talonarios coinciden con la forma real de trabajo.", styles["BodyX"]))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawString(0.65 * inch, 0.35 * inch, "Informe modulo de cobranza")
        canvas.drawRightString(7.85 * inch, 0.35 * inch, f"Pagina {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build():
    make_flow()
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    build()
