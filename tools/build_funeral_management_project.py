from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "proyecto-sistema-gestion-servicios-funebres.docx"
PDF_PATH = OUT / "proyecto-sistema-gestion-servicios-funebres.pdf"
IMG_MODULES = OUT / "grafico-modulos-gestion-funebre.png"
IMG_FLOW = OUT / "grafico-flujo-servicio-funebre.png"


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(20, 32, 48)
LIGHT = "F4F6F9"
MID = "D9EAF7"


def font(size=28, bold=False):
    try:
        return ImageFont.truetype("arialbd.ttf" if bold else "arial.ttf", size)
    except Exception:
        return ImageFont.load_default()


def draw_box(draw, xy, text, fill, outline="#2E74B5", text_fill="#142030", title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    words = text.split()
    lines, line = [], ""
    max_chars = 18 if title else 24
    for word in words:
        if len((line + " " + word).strip()) > max_chars:
            lines.append(line)
            line = word
        else:
            line = (line + " " + word).strip()
    if line:
        lines.append(line)
    f = font(26 if title else 22, bold=title)
    total_h = len(lines) * 28
    y = y1 + ((y2 - y1) - total_h) / 2
    for ln in lines:
        bbox = draw.textbbox((0, 0), ln, font=f)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), ln, font=f, fill=text_fill)
        y += 28


def make_graphics():
    img = Image.new("RGB", (1400, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 45), "Sistema de gestion integral", font=font(42, True), fill="#142030")
    d.text((60, 100), "Modulos conectados para administracion, servicios, caja y cobranza", font=font(24), fill="#4B5870")
    center = (540, 380, 860, 560)
    draw_box(d, center, "Gestion central", "#E8EEF5", "#1F4D78", title=True)
    boxes = [
        ((70, 210, 360, 330), "Afiliados y planes"),
        ((70, 420, 360, 540), "Cobranza mensual"),
        ((70, 630, 360, 750), "Caja y turnos"),
        ((520, 640, 880, 760), "Reportes y control"),
        ((1040, 630, 1330, 750), "Stock e inventario"),
        ((1040, 420, 1330, 540), "Facturacion"),
        ((1040, 210, 1330, 330), "Remitos y servicios"),
        ((520, 175, 880, 295), "Usuarios y roles"),
    ]
    for xy, text in boxes:
        draw_box(d, xy, text, "#F4F6F9")
        x1, y1, x2, y2 = xy
        d.line(((x1 + x2) / 2, (y1 + y2) / 2, 700, 470), fill="#9AB8D4", width=4)
    img.save(IMG_MODULES)

    img = Image.new("RGB", (1400, 620), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 45), "Flujo operativo propuesto", font=font(42, True), fill="#142030")
    steps = [
        "Carga del servicio",
        "Seleccion de mejoras",
        "Remito y documentacion",
        "Facturacion",
        "Caja / cobro / saldo",
        "Reporte y control",
    ]
    x = 55
    for i, step in enumerate(steps):
        draw_box(d, (x, 210, x + 195, 345), step, "#E8EEF5" if i % 2 == 0 else "#F4F6F9")
        if i < len(steps) - 1:
            d.line((x + 195, 277, x + 232, 277), fill="#2E74B5", width=5)
            d.polygon([(x + 232, 277), (x + 218, 267), (x + 218, 287)], fill="#2E74B5")
        x += 220
    d.text((60, 430), "Cada paso conserva historial, usuario responsable, comprobantes y estado administrativo.", font=font(26), fill="#4B5870")
    img.save(IMG_FLOW)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK)]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proyecto de Sistema de Gestion Integral")
    r.font.name = "Calibri"
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = INK
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Empresa de servicios funebres, afiliados, cobranza y administracion operativa")
    r.font.size = Pt(14)
    r.font.color.rgb = DARK
    doc.add_paragraph()
    doc.add_picture(str(IMG_MODULES), width=Inches(6.5))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Documento de presentacion funcional y comercial").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Version preliminar para validacion del alcance")
    doc.add_page_break()


def build():
    make_graphics()
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "El proyecto propone desarrollar una plataforma de gestion integral para una empresa dedicada a la venta, "
        "administracion y atencion de servicios funebres. La solucion parte del modulo de cobranza ya trabajado y "
        "lo expande hacia un sistema centralizado para afiliados, servicios, remitos, caja, facturacion, stock, "
        "documentacion, alertas y reportes."
    )
    doc.add_paragraph(
        "El objetivo es reemplazar planillas dispersas y registros manuales por una herramienta unica, con usuarios "
        "y roles, trazabilidad de acciones, control financiero por turno y seguimiento operativo de cada servicio."
    )

    doc.add_heading("2. Objetivo del sistema", level=1)
    add_bullets(doc, [
        "Centralizar la informacion de afiliados, clientes, cobranzas, servicios y convenios.",
        "Controlar la cobranza mensual de seguros, recibos, tickets, saldos y rendiciones.",
        "Gestionar servicios funebres desde la carga inicial hasta remito, facturacion y cierre administrativo.",
        "Registrar caja, movimientos, gastos, depositos bancarios y traspaso de turnos.",
        "Controlar documentacion requerida por cada convenio antes de facturar y presentar.",
        "Administrar stock de ataudes, insumos funerarios y articulos operativos.",
        "Generar reportes para supervision, administracion y toma de decisiones.",
    ])

    doc.add_heading("3. Alcance general por modulos", level=1)
    add_table(doc, ["Modulo", "Funcion principal", "Resultado esperado"], [
        ["Afiliados y planes", "Base unica de clientes, polizas, planes, valores y dependencia/cobrador.", "Cartera ordenada y actualizable."],
        ["Cobranza", "Tickets, recibos, medios de pago, novedades, totales y rendicion.", "Control mensual de ingresos y pendientes."],
        ["Servicios funebres", "Carga de servicio basico, fallecido, solicitante, cobertura y mejoras.", "Orden operativa completa por servicio."],
        ["Remitos", "Documento del servicio prestado con items, importes y observaciones.", "Constancia formal para control y facturacion."],
        ["Facturacion", "Remitos pendientes, datos fiscales, comprobantes y estado de factura.", "Seguimiento administrativo hasta cierre."],
        ["Caja y turnos", "Ingresos, egresos, depositos, saldos y reporte firmado por turno.", "Control de caja y continuidad entre turnos."],
        ["Documentacion", "Checklist por convenio, estado de papeles y archivos.", "Saber que falta para facturar/presentar."],
        ["Stock", "Ataudes, urnas, cafe, azucar, te, limpieza e insumos.", "Control de existencias, consumos y alertas."],
        ["Usuarios y roles", "Permisos por perfil y registro de acciones.", "Seguridad y trazabilidad."],
    ], widths=[1.55, 2.8, 2.15])

    doc.add_heading("4. Modulo de cobranza ya iniciado", level=1)
    doc.add_paragraph(
        "El modulo de cobranza es la base operativa inicial del sistema. Permite administrar la cartera de afiliados, "
        "preparar la cobranza mensual y controlar lo cobrado por tickets, recibos, efectivo y transferencias."
    )
    add_table(doc, ["Componente", "Detalle"], [
        ["Base de afiliados", "Nombre, poliza, plan, valor, telefono, direccion, dependencia, cobrador y cantidad de tickets."],
        ["Seleccion mensual", "Tildar afiliados uno por uno o todos; cargar los seleccionados al mes de cobranza."],
        ["Aumento de valores", "Aplicar porcentaje de aumento a la planilla mensual, con opcion de deshacer el ultimo aumento."],
        ["Cobranza por tickets", "Busqueda por poliza y plan, tickets disponibles, cantidad cobrada y medio E/T."],
        ["Transferencias", "Registro de numero de comprobante y datos asociados cuando el pago es por transferencia."],
        ["Recibos", "Carga de recibos sin tickets con numero de recibo, plan, mes, monto y medio de pago."],
        ["Novedades", "Observaciones libres sobre afiliados, domicilios, telefonos, visitas y pendientes."],
        ["WhatsApp cobrador", "Mensaje al cobrador con datos del afiliado y cantidad de tickets pendientes."],
        ["Totales", "Tickets recibidos, cobrados, pendientes, recibos, efectivo y transferencias por plan."],
        ["Rendicion", "Total efectivo, total transferencia, comision 12%, total a rendir, rendido y diferencia."],
    ], widths=[1.8, 4.7])

    doc.add_heading("5. Gestion de servicios funebres", level=1)
    doc.add_paragraph(
        "La empresa vende y atiende distintos tipos de servicios funebres. Cada servicio puede originarse por seguro, "
        "pago particular, efectivo, convenio con sindicato, obra social u otra entidad."
    )
    doc.add_picture(str(IMG_FLOW), width=Inches(6.5))
    doc.add_heading("Servicio basico y mejoras", level=2)
    doc.add_paragraph(
        "El sistema permitira cargar un servicio basico y agregar mejoras mediante casillas de seleccion. Cada mejora "
        "tendra costos preestablecidos desde un catalogo administrable."
    )
    add_table(doc, ["Item", "Como funcionaria"], [
        ["Servicio basico", "Precio base del servicio, cobertura y datos principales del fallecido/solicitante."],
        ["Mejoras", "Coche de acompanamiento, buffet, arreglos florales, ataud superior, traslados, sala especial y otros."],
        ["Costos predefinidos", "Cada mejora trae su valor automaticamente, con cantidad y subtotal."],
        ["Total servicio", "Suma del servicio basico, mejoras, descuentos, cobertura y saldo pendiente."],
    ], widths=[1.8, 4.7])

    doc.add_heading("6. Remitos", level=1)
    doc.add_paragraph(
        "Una vez cargado el servicio, se generara un remito con los datos necesarios para dejar constancia del servicio "
        "prestado y preparar la facturacion."
    )
    add_bullets(doc, [
        "Numero de remito, fecha y usuario que lo genero.",
        "Datos del fallecido, solicitante, titular o afiliado relacionado.",
        "Tipo de cobertura o pago: seguro, particular, sindicato, obra social o convenio.",
        "Servicio basico, mejoras seleccionadas, cantidades, importes y total.",
        "Saldo pendiente y posible fecha de pago cuando corresponda.",
        "Observaciones y estado del remito: borrador, confirmado, anulado o facturado.",
    ])

    doc.add_heading("7. Facturacion y documentacion por convenio", level=1)
    doc.add_paragraph(
        "Cuando se genera un remito, la informacion pasa al modulo de facturacion como pendiente. Antes de facturar, "
        "el sistema permitira revisar la documentacion requerida por el convenio correspondiente."
    )
    add_table(doc, ["Etapa", "Control"], [
        ["Remito pendiente", "El sistema muestra remitos aun no facturados, con importes y tipo de cobertura."],
        ["Checklist documental", "Cada convenio trae su lista de documentos requeridos."],
        ["Estado documental", "Pendiente, parcial, completo, observado o no corresponde."],
        ["Factura emitida", "Carga de tipo, punto de venta, numero, fecha, importe, CAE/comprobante y observaciones."],
        ["Presentacion", "Registro de presentacion al convenio, fecha y estado de cobro."],
    ], widths=[1.8, 4.7])

    doc.add_heading("8. Saldos particulares y alertas de cobro", level=1)
    doc.add_paragraph(
        "Para servicios particulares con saldo pendiente, el remito permitira cargar una fecha probable o comprometida "
        "de pago. El sistema emitira alertas 2 o 3 dias antes del vencimiento para que administracion contacte al deudor."
    )
    add_bullets(doc, [
        "Registro de saldo pendiente, deudor, telefono y fecha comprometida.",
        "Panel de alertas: proximo a vencer, vence hoy y vencido.",
        "Carga de novedades libres sobre cada comunicacion.",
        "Historial de llamados, mensajes, acuerdos, prorroga o pago recibido.",
    ])

    doc.add_heading("9. Caja, movimientos y turnos", level=1)
    doc.add_paragraph(
        "El sistema administrara caja por turno, incluyendo cobranzas, gastos, depositos bancarios, saldos y traspaso "
        "de informacion al turno siguiente."
    )
    add_table(doc, ["Movimiento", "Ejemplos"], [
        ["Cobranzas", "Servicios, mejoras, saldos particulares, recibos, cuotas, convenios y transferencias."],
        ["Gastos", "Proveedores, combustible, traslados, flores, buffet, cementerio, limpieza e insumos."],
        ["Depositos bancarios", "Retiro de efectivo de caja, banco/cuenta, monto, comprobante y usuario responsable."],
        ["Cierre de turno", "Caja inicial, ingresos, egresos, saldo esperado, contado, diferencia y observaciones."],
        ["Reporte firmado", "Impresion/PDF para firma de turno saliente y entrante."],
    ], widths=[1.8, 4.7])

    doc.add_heading("10. Stock e inventario", level=1)
    doc.add_paragraph(
        "La plataforma incluira control de stock para articulos funerarios y articulos operativos de consumo diario."
    )
    add_table(doc, ["Tipo de stock", "Articulos", "Control"], [
        ["Funerario", "Ataudes, urnas, placas, velas, mortajas y elementos del servicio.", "Stock actual, minimo, reservado, usado o dado de baja."],
        ["Operativo", "Cafe, azucar, te, vasos, servilletas, limpieza e insumos administrativos.", "Ingresos, consumos, ajustes y alertas de reposicion."],
        ["Conexion con remitos", "Ataud y articulos usados en cada servicio.", "Descuento automatico y costo asociado al servicio."],
    ], widths=[1.4, 2.6, 2.5])

    doc.add_heading("11. Usuarios y roles", level=1)
    add_table(doc, ["Rol", "Permisos principales"], [
        ["Administrador", "Acceso total, configuracion, usuarios, precios, anulaciones y reportes."],
        ["Administracion", "Afiliados, cobranza, facturacion, documentacion, caja y reportes administrativos."],
        ["Cobrador", "Cartera asignada, tickets, recibos, novedades y rendicion propia."],
        ["Operador de servicios", "Carga de servicios, mejoras, remitos, documentacion y novedades operativas."],
        ["Supervisor", "Control de cobradores, servicios, caja, turnos, diferencias y autorizaciones."],
    ], widths=[1.8, 4.7])

    doc.add_heading("12. Roadmap sugerido", level=1)
    add_table(doc, ["Etapa", "Objetivo", "Resultado"], [
        ["1. Consolidacion cobranza", "Ajustar modulo actual con pruebas reales.", "Base estable de afiliados, tickets, recibos y rendicion."],
        ["2. Servicios y remitos", "Crear catalogo de mejoras y remito del servicio.", "Operacion funeraria registrada en sistema."],
        ["3. Caja y turnos", "Registrar movimientos y cierre firmado.", "Control diario de dinero e informacion transferida."],
        ["4. Facturacion/documentacion", "Conectar remitos con facturacion y checklist por convenio.", "Menos demoras y pendientes visibles."],
        ["5. Stock y reportes", "Controlar insumos, ataudes, consumos y tableros.", "Gestion integral con indicadores."],
        ["6. App movil / nube", "Usuarios, permisos, base online y uso desde telefono.", "Escalabilidad y trabajo multiusuario."],
    ], widths=[1.5, 2.6, 2.4])

    doc.add_heading("13. Beneficios esperados", level=1)
    add_bullets(doc, [
        "Menos errores por planillas manuales y datos duplicados.",
        "Mayor control sobre cobranzas, saldos, caja y rendiciones.",
        "Seguimiento completo de servicios funebres desde venta hasta facturacion.",
        "Alertas para saldos pendientes, documentacion faltante y stock bajo.",
        "Reportes por mes, convenio, cobrador, servicio, caja y turno.",
        "Base preparada para crecer hacia una aplicacion web/movil multiusuario.",
    ])

    doc.add_heading("14. Cierre", level=1)
    doc.add_paragraph(
        "La propuesta busca transformar la operatoria actual en un sistema de gestion integral, comenzando por la "
        "cobranza ya desarrollada y ampliando progresivamente hacia el control completo de la empresa. El enfoque "
        "recomendado es avanzar por etapas, probar con casos reales y ajustar cada modulo segun el uso diario."
    )

    doc.save(DOCX_PATH)
    build_pdf()
    return DOCX_PATH


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CoverTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=25,
        leading=30,
        textColor=colors.HexColor("#142030"),
        alignment=TA_CENTER,
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name="CoverSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#1F4D78"),
        alignment=TA_CENTER,
        spaceAfter=18,
    ))
    styles.add(ParagraphStyle(
        name="H1x",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=19,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=14,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="H2x",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=10,
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="Bodyx",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.6,
        leading=13.5,
        spaceAfter=6,
        textColor=colors.HexColor("#202A3A"),
    ))
    styles.add(ParagraphStyle(
        name="Bulletx",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.4,
        leading=13,
        leftIndent=14,
        firstLineIndent=-8,
        bulletIndent=0,
        spaceAfter=4,
        textColor=colors.HexColor("#202A3A"),
    ))
    styles.add(ParagraphStyle(
        name="Smallx",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#4B5870"),
    ))
    return styles


def p(text, style):
    return Paragraph(text, style)


def pdf_table(data, widths, styles):
    converted = [[p(str(cell), styles["Smallx"]) for cell in row] for row in data]
    tbl = Table(converted, colWidths=[w * inch for w in widths], repeatRows=1, hAlign="CENTER")
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#142030")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return tbl


def pdf_bullets(story, styles, items):
    for item in items:
        story.append(Paragraph(f"• {item}", styles["Bulletx"]))


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.55 * inch,
    )
    story = []
    story.append(Spacer(1, 0.3 * inch))
    story.append(p("Proyecto de Sistema de Gestion Integral", styles["CoverTitle"]))
    story.append(p("Empresa de servicios funebres, afiliados, cobranza y administracion operativa", styles["CoverSubtitle"]))
    story.append(PdfImage(str(IMG_MODULES), width=6.6 * inch, height=3.87 * inch))
    story.append(Spacer(1, 0.25 * inch))
    story.append(p("<b>Documento de presentacion funcional y comercial</b>", styles["CoverSubtitle"]))
    story.append(p("Version preliminar para validacion del alcance", styles["Smallx"]))
    story.append(PageBreak())

    story.append(p("1. Resumen ejecutivo", styles["H1x"]))
    story.append(p("El proyecto propone desarrollar una plataforma de gestion integral para una empresa dedicada a la venta, administracion y atencion de servicios funebres. La solucion parte del modulo de cobranza ya trabajado y lo expande hacia un sistema centralizado para afiliados, servicios, remitos, caja, facturacion, stock, documentacion, alertas y reportes.", styles["Bodyx"]))
    story.append(p("El objetivo es reemplazar planillas dispersas y registros manuales por una herramienta unica, con usuarios y roles, trazabilidad de acciones, control financiero por turno y seguimiento operativo de cada servicio.", styles["Bodyx"]))

    story.append(p("2. Objetivo del sistema", styles["H1x"]))
    pdf_bullets(story, styles, [
        "Centralizar afiliados, clientes, cobranzas, servicios y convenios.",
        "Controlar cobranza mensual de seguros, recibos, tickets, saldos y rendiciones.",
        "Gestionar servicios funebres desde carga inicial hasta remito, facturacion y cierre administrativo.",
        "Registrar caja, movimientos, gastos, depositos bancarios y traspaso de turnos.",
        "Controlar documentacion requerida por convenio antes de facturar y presentar.",
        "Administrar stock de ataudes, insumos funerarios y articulos operativos.",
    ])

    story.append(p("3. Alcance general por modulos", styles["H1x"]))
    story.append(pdf_table([
        ["Modulo", "Funcion principal", "Resultado esperado"],
        ["Afiliados y planes", "Base unica de clientes, polizas, planes, valores y dependencia/cobrador.", "Cartera ordenada y actualizable."],
        ["Cobranza", "Tickets, recibos, medios de pago, novedades, totales y rendicion.", "Control mensual de ingresos y pendientes."],
        ["Servicios funebres", "Carga de servicio basico, fallecido, solicitante, cobertura y mejoras.", "Orden operativa completa por servicio."],
        ["Remitos", "Documento del servicio prestado con items, importes y observaciones.", "Constancia formal para control y facturacion."],
        ["Facturacion", "Remitos pendientes, datos fiscales, comprobantes y estado de factura.", "Seguimiento administrativo hasta cierre."],
        ["Caja y turnos", "Ingresos, egresos, depositos, saldos y reporte firmado por turno.", "Control de caja y continuidad entre turnos."],
        ["Documentacion", "Checklist por convenio, estado de papeles y archivos.", "Saber que falta para facturar/presentar."],
        ["Stock", "Ataudes, urnas, cafe, azucar, te, limpieza e insumos.", "Control de existencias, consumos y alertas."],
    ], [1.45, 3.0, 2.0], styles))

    story.append(PageBreak())
    story.append(p("4. Modulo de cobranza ya iniciado", styles["H1x"]))
    story.append(p("El modulo de cobranza es la base operativa inicial del sistema. Permite administrar la cartera de afiliados, preparar la cobranza mensual y controlar lo cobrado por tickets, recibos, efectivo y transferencias.", styles["Bodyx"]))
    story.append(pdf_table([
        ["Componente", "Detalle"],
        ["Base de afiliados", "Nombre, poliza, plan, valor, telefono, direccion, dependencia, cobrador y cantidad de tickets."],
        ["Seleccion mensual", "Tildar afiliados uno por uno o todos; cargar los seleccionados al mes de cobranza."],
        ["Aumento de valores", "Aplicar porcentaje de aumento a la planilla mensual, con opcion de deshacer el ultimo aumento."],
        ["Cobranza por tickets", "Busqueda por poliza y plan, tickets disponibles, cantidad cobrada y medio E/T."],
        ["Transferencias", "Registro de numero de comprobante y datos asociados cuando el pago es por transferencia."],
        ["Recibos", "Carga de recibos sin tickets con numero de recibo, plan, mes, monto y medio de pago."],
        ["Novedades", "Observaciones libres sobre afiliados, domicilios, telefonos, visitas y pendientes."],
        ["WhatsApp cobrador", "Mensaje al cobrador con datos del afiliado y cantidad de tickets pendientes."],
        ["Totales y rendicion", "Tickets recibidos, cobrados, pendientes, efectivo, transferencias, comision 12% y total a rendir."],
    ], [1.65, 4.8], styles))

    story.append(p("5. Gestion de servicios funebres", styles["H1x"]))
    story.append(p("La empresa vende y atiende distintos tipos de servicios funebres. Cada servicio puede originarse por seguro, pago particular, efectivo, convenio con sindicato, obra social u otra entidad.", styles["Bodyx"]))
    story.append(PdfImage(str(IMG_FLOW), width=6.6 * inch, height=2.92 * inch))
    story.append(p("Servicio basico y mejoras", styles["H2x"]))
    story.append(p("El sistema permitira cargar un servicio basico y agregar mejoras mediante casillas de seleccion. Cada mejora tendra costos preestablecidos desde un catalogo administrable.", styles["Bodyx"]))
    story.append(pdf_table([
        ["Item", "Como funcionaria"],
        ["Servicio basico", "Precio base del servicio, cobertura y datos principales del fallecido/solicitante."],
        ["Mejoras", "Coche de acompanamiento, buffet, arreglos florales, ataud superior, traslados, sala especial y otros."],
        ["Costos predefinidos", "Cada mejora trae su valor automaticamente, con cantidad y subtotal."],
        ["Total servicio", "Suma del servicio basico, mejoras, descuentos, cobertura y saldo pendiente."],
    ], [1.65, 4.8], styles))

    story.append(PageBreak())
    story.append(p("6. Remitos", styles["H1x"]))
    pdf_bullets(story, styles, [
        "Numero de remito, fecha y usuario que lo genero.",
        "Datos del fallecido, solicitante, titular o afiliado relacionado.",
        "Tipo de cobertura o pago: seguro, particular, sindicato, obra social o convenio.",
        "Servicio basico, mejoras seleccionadas, cantidades, importes y total.",
        "Saldo pendiente y posible fecha de pago cuando corresponda.",
        "Observaciones y estado: borrador, confirmado, anulado o facturado.",
    ])

    story.append(p("7. Facturacion y documentacion por convenio", styles["H1x"]))
    story.append(pdf_table([
        ["Etapa", "Control"],
        ["Remito pendiente", "El sistema muestra remitos aun no facturados, con importes y tipo de cobertura."],
        ["Checklist documental", "Cada convenio trae su lista de documentos requeridos."],
        ["Estado documental", "Pendiente, parcial, completo, observado o no corresponde."],
        ["Factura emitida", "Carga de tipo, punto de venta, numero, fecha, importe, CAE/comprobante y observaciones."],
        ["Presentacion", "Registro de presentacion al convenio, fecha y estado de cobro."],
    ], [1.65, 4.8], styles))

    story.append(p("8. Saldos particulares y alertas de cobro", styles["H1x"]))
    pdf_bullets(story, styles, [
        "Registro de saldo pendiente, deudor, telefono y fecha comprometida.",
        "Panel de alertas: proximo a vencer, vence hoy y vencido.",
        "Carga de novedades libres sobre cada comunicacion.",
        "Historial de llamados, mensajes, acuerdos, prorroga o pago recibido.",
    ])

    story.append(p("9. Caja, movimientos y turnos", styles["H1x"]))
    story.append(pdf_table([
        ["Movimiento", "Ejemplos"],
        ["Cobranzas", "Servicios, mejoras, saldos particulares, recibos, cuotas, convenios y transferencias."],
        ["Gastos", "Proveedores, combustible, traslados, flores, buffet, cementerio, limpieza e insumos."],
        ["Depositos bancarios", "Retiro de efectivo de caja, banco/cuenta, monto, comprobante y usuario responsable."],
        ["Cierre de turno", "Caja inicial, ingresos, egresos, saldo esperado, contado, diferencia y observaciones."],
        ["Reporte firmado", "Impresion/PDF para firma de turno saliente y entrante."],
    ], [1.65, 4.8], styles))

    story.append(PageBreak())
    story.append(p("10. Stock e inventario", styles["H1x"]))
    story.append(pdf_table([
        ["Tipo", "Articulos", "Control"],
        ["Funerario", "Ataudes, urnas, placas, velas, mortajas y elementos del servicio.", "Stock actual, minimo, reservado, usado o dado de baja."],
        ["Operativo", "Cafe, azucar, te, vasos, servilletas, limpieza e insumos administrativos.", "Ingresos, consumos, ajustes y alertas de reposicion."],
        ["Conexion con remitos", "Ataud y articulos usados en cada servicio.", "Descuento automatico y costo asociado al servicio."],
    ], [1.4, 2.65, 2.4], styles))

    story.append(p("11. Usuarios y roles", styles["H1x"]))
    story.append(pdf_table([
        ["Rol", "Permisos principales"],
        ["Administrador", "Acceso total, configuracion, usuarios, precios, anulaciones y reportes."],
        ["Administracion", "Afiliados, cobranza, facturacion, documentacion, caja y reportes administrativos."],
        ["Cobrador", "Cartera asignada, tickets, recibos, novedades y rendicion propia."],
        ["Operador de servicios", "Carga de servicios, mejoras, remitos, documentacion y novedades operativas."],
        ["Supervisor", "Control de cobradores, servicios, caja, turnos, diferencias y autorizaciones."],
    ], [1.65, 4.8], styles))

    story.append(p("12. Roadmap sugerido", styles["H1x"]))
    story.append(pdf_table([
        ["Etapa", "Objetivo", "Resultado"],
        ["1. Consolidacion cobranza", "Ajustar modulo actual con pruebas reales.", "Base estable de afiliados, tickets, recibos y rendicion."],
        ["2. Servicios y remitos", "Crear catalogo de mejoras y remito del servicio.", "Operacion funeraria registrada en sistema."],
        ["3. Caja y turnos", "Registrar movimientos y cierre firmado.", "Control diario de dinero e informacion transferida."],
        ["4. Facturacion/documentacion", "Conectar remitos con facturacion y checklist por convenio.", "Menos demoras y pendientes visibles."],
        ["5. Stock y reportes", "Controlar insumos, ataudes, consumos y tableros.", "Gestion integral con indicadores."],
        ["6. App movil / nube", "Usuarios, permisos, base online y uso desde telefono.", "Escalabilidad y trabajo multiusuario."],
    ], [1.45, 2.55, 2.45], styles))

    story.append(p("13. Beneficios esperados", styles["H1x"]))
    pdf_bullets(story, styles, [
        "Menos errores por planillas manuales y datos duplicados.",
        "Mayor control sobre cobranzas, saldos, caja y rendiciones.",
        "Seguimiento completo de servicios funebres desde venta hasta facturacion.",
        "Alertas para saldos pendientes, documentacion faltante y stock bajo.",
        "Reportes por mes, convenio, cobrador, servicio, caja y turno.",
        "Base preparada para crecer hacia una aplicacion web/movil multiusuario.",
    ])
    story.append(p("14. Cierre", styles["H1x"]))
    story.append(p("La propuesta busca transformar la operatoria actual en un sistema de gestion integral, comenzando por la cobranza ya desarrollada y ampliando progresivamente hacia el control completo de la empresa. El enfoque recomendado es avanzar por etapas, probar con casos reales y ajustar cada modulo segun el uso diario.", styles["Bodyx"]))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawString(0.65 * inch, 0.35 * inch, "Proyecto sistema de gestion integral")
        canvas.drawRightString(7.85 * inch, 0.35 * inch, f"Pagina {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    print(build())
